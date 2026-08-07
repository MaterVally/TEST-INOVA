"""
DOCX preprocessing pipeline.

Enterprise Compliance Intelligence Platform

Purpose
-------
Process Microsoft Word documents into the same format produced by
PdfChunking.

Output
------
texts  : List[str]
images : List[dict]

The interface intentionally matches PdfChunking so MMKGBuilder can
switch processors based only on file extension.
"""

from __future__ import annotations

import os
import shutil
import zipfile

from docx import Document

from ..utils.base import logger


class DocxChunking:

    def __init__(
        self,
        docx_path: str,
        working_dir: str,
    ):

        self.docx_path = docx_path
        self.working_dir = working_dir

        self.images_dir = os.path.join(
            working_dir,
            "images"
        )

        os.makedirs(
            self.images_dir,
            exist_ok=True
        )

    # ---------------------------------------------------------
    # Public Entry
    # ---------------------------------------------------------

    async def process(
        self
    ) -> tuple[list[str], list[dict]]:

        logger.info("📄 Processing DOCX document...")

        texts = self._extract_text()

        images = self._extract_images()

        logger.info(
            f"✅ DOCX Parsed "
            f"({len(texts)} text blocks, "
            f"{len(images)} images)"
        )

        return texts, images

    # ---------------------------------------------------------
    # Text Extraction
    # ---------------------------------------------------------

    def _extract_text(self) -> list[str]:

        document = Document(
            self.docx_path
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                paragraphs.append(text)

        # Also read tables

        for table in document.tables:

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                if any(cells):

                    paragraphs.append(
                        " | ".join(cells)
                    )

        return paragraphs

    # ---------------------------------------------------------
    # Embedded Image Extraction
    # ---------------------------------------------------------

    def _extract_images(self) -> list[dict]:

        extracted = []

        tmp_dir = os.path.join(
            self.working_dir,
            "_docx_tmp"
        )

        if os.path.exists(tmp_dir):

            shutil.rmtree(tmp_dir)

        os.makedirs(
            tmp_dir,
            exist_ok=True
        )

        with zipfile.ZipFile(
            self.docx_path,
            "r"
        ) as archive:

            archive.extractall(tmp_dir)

        media_dir = os.path.join(
            tmp_dir,
            "word",
            "media"
        )

        if os.path.exists(media_dir):

            for filename in os.listdir(media_dir):

                source = os.path.join(
                    media_dir,
                    filename
                )

                destination = os.path.join(
                    self.images_dir,
                    filename
                )

                shutil.copy2(
                    source,
                    destination
                )

                extracted.append(

                    {
                        "image_name": filename,

                        "image_path": destination,

                        "description": ""

                    }

                )

        shutil.rmtree(
            tmp_dir,
            ignore_errors=True
        )

        return extracted
